"""
style_formatter.py

Word字体格式处理模块

功能:

1. 设置Normal默认样式
2. 修改文档所有Run字体
3. 统一中文字体
4. 统一英文字体
5. 统一字号
6. 清除原有字符格式
7. 小括号及括号内容设置楷体小三


说明:

本模块只负责字体。

不负责:

1. 标题识别
2. 标题格式
3. 段落格式
4. 页面格式


依赖:

python-docx

"""


import re


from docx import Document

from docx.text.run import Run

from docx.shared import Pt

from docx.oxml.ns import qn


from config import (

    BODY_FONT_CN,

    BODY_FONT_EN,

    BODY_FONT_SIZE

)


# =====================================================
# 括号字体配置
# =====================================================


BRACKET_FONT_CN = (

    "楷体_GB2312"

)


BRACKET_FONT_EN = (

    "Times New Roman"

)


# 小三

BRACKET_FONT_SIZE = 15


# 中文小括号 + 英文小括号

BRACKET_PATTERN = re.compile(

    r"（.*?）|\(.*?\)"

)


# =====================================================
# 清除字符特殊格式
# =====================================================


def clear_run_format(
        run: Run
):

    run.bold = False

    run.italic = False

    run.underline = False

    run.font.strike = False

    run.font.double_strike = False

    run.font.superscript = False

    run.font.subscript = False

    run.font.color.rgb = None

    run.font.highlight_color = None


# =====================================================
# 设置普通字体
# =====================================================


def set_run_font(
        run: Run
):

    clear_run_format(
        run
    )

    rPr = (
        run._element
        .get_or_add_rPr()
    )

    rFonts = (
        rPr
        .get_or_add_rFonts()
    )

    rFonts.set(
        qn("w:eastAsia"),
        BODY_FONT_CN
    )

    rFonts.set(
        qn("w:ascii"),
        BODY_FONT_EN
    )

    rFonts.set(
        qn("w:hAnsi"),
        BODY_FONT_EN
    )

    run.font.name = BODY_FONT_EN

    run.font.size = Pt(
        BODY_FONT_SIZE
    )


# =====================================================
# 设置括号字体
# =====================================================


def set_bracket_font(
        run: Run
):

    clear_run_format(
        run
    )

    rPr = (
        run._element
        .get_or_add_rPr()
    )

    rFonts = (
        rPr
        .get_or_add_rFonts()
    )

    rFonts.set(
        qn("w:eastAsia"),
        BRACKET_FONT_CN
    )

    rFonts.set(
        qn("w:ascii"),
        BRACKET_FONT_EN
    )

    rFonts.set(
        qn("w:hAnsi"),
        BRACKET_FONT_EN
    )

    run.font.name = BRACKET_FONT_EN

    run.font.size = Pt(
        BRACKET_FONT_SIZE
    )

    # 强制取消加粗

    run.bold = False


# =====================================================
# 处理括号文字
# =====================================================


def format_bracket_runs(
        paragraph
):
    """
    查找:

        （xxx）

        (xxx)


    并设置:

        楷体
        小三
        不加粗

    """

    for run in list(
            paragraph.runs
    ):

        text = run.text

        if not text:

            continue

        matches = list(
            BRACKET_PATTERN.finditer(
                text
            )
        )

        if not matches:

            continue

        # 保存原Run位置

        parent = (
            run._element
            .getparent()
        )

        index = (
            parent
            .index(run._element)
        )

        # 删除原Run

        parent.remove(
            run._element
        )

        last = 0

        for match in matches:

            start, end = match.span()

            # 前面的普通文字

            if start > last:

                new_run = (
                    paragraph
                    .add_run(
                        text[last:start]
                    )
                )

                set_run_font(
                    new_run
                )

            # 括号内容

            bracket_run = (
                paragraph
                .add_run(
                    text[start:end]
                )
            )

            set_bracket_font(
                bracket_run
            )

            last = end

        # 后面的普通文字

        if last < len(text):

            new_run = (
                paragraph
                .add_run(
                    text[last:]
                )
            )

            set_run_font(
                new_run
            )


# =====================================================
# 设置Normal样式
# =====================================================


def set_normal_style(
        document: Document
):

    style = (
        document.styles["Normal"]
    )

    style.font.name = BODY_FONT_EN

    style.font.size = Pt(
        BODY_FONT_SIZE
    )

    rPr = (
        style._element
        .get_or_add_rPr()
    )

    rFonts = (
        rPr
        .get_or_add_rFonts()
    )

    rFonts.set(
        qn("w:eastAsia"),
        BODY_FONT_CN
    )

    rFonts.set(
        qn("w:ascii"),
        BODY_FONT_EN
    )

    rFonts.set(
        qn("w:hAnsi"),
        BODY_FONT_EN
    )


# =====================================================
# 遍历所有文字
# =====================================================


def format_all_runs(
        document: Document
):

    for paragraph in document.paragraphs:

        for run in paragraph.runs:

            set_run_font(
                run
            )

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    for run in paragraph.runs:

                        set_run_font(
                            run
                        )


# =====================================================
# 处理括号
# =====================================================


def format_all_brackets(
        document: Document
):

    for paragraph in document.paragraphs:

        format_bracket_runs(
            paragraph
        )

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                for paragraph in cell.paragraphs:

                    format_bracket_runs(
                        paragraph
                    )


# =====================================================
# 总入口
# =====================================================


def format_document_style(
        document: Document
):

    set_normal_style(
        document
    )

    # 统一正文

    format_all_runs(
        document
    )

    # 最后覆盖括号

    format_all_brackets(
        document
    )

    return document


# =====================================================
# 测试
# =====================================================

if __name__ == "__main__":

    from config import INPUT_FILE

    doc = Document(
        INPUT_FILE
    )

    format_document_style(
        doc
    )

    doc.save(
        "style_test.docx"
    )

    print(
        "字体格式处理完成"
    )
