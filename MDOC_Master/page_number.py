"""
page_number.py

Word页码处理模块

功能:

2. 插入自动页码
3. 设置页码字体
4. 设置页码格式
5. 设置页码位置


参数来源:

    config.py


依赖:

    python-docx
"""


from docx import Document

from docx.shared import Pt

from docx.enum.text import (
    WD_ALIGN_PARAGRAPH
)

from docx.oxml import OxmlElement

from docx.oxml.ns import qn


from config import (

    PAGE_NUMBER_FONT,

    PAGE_NUMBER_SIZE,

    PAGE_NUMBER_FORMAT,

    PAGE_NUMBER_ALIGNMENT

)


# =====================================================
# 设置页码字体
# =====================================================


def set_page_number_font(
        run
):
    """
    设置页码字体
    """

    rPr = (
        run._element
        .get_or_add_rPr()
    )

    rFonts = (
        rPr
        .get_or_add_rFonts()
    )

    # 中文字体

    rFonts.set(
        qn("w:eastAsia"),
        PAGE_NUMBER_FONT
    )

    # 英文数字字体

    rFonts.set(
        qn("w:ascii"),
        "Times New Roman"
    )

    rFonts.set(
        qn("w:hAnsi"),
        "Times New Roman"
    )

    rFonts.set(
        qn("w:cs"),
        "Times New Roman"
    )

    run.font.size = Pt(
        PAGE_NUMBER_SIZE
    )


# =====================================================
# 清空页脚
# =====================================================


def clear_footer(
        footer
):
    """
    清除页脚内容

    不破坏pPr
    """

    for paragraph in footer.paragraphs:

        for run in paragraph.runs:

            run._element.getparent().remove(
                run._element
            )


# =====================================================
# 创建Word PAGE域
# =====================================================


def add_page_field(
        paragraph
):
    """
    创建标准PAGE域


    XML结构:

        begin

        instrText PAGE

        separate

        end


    每个Field元素独立Run

    """

    # -----------------------------
    # begin
    # -----------------------------

    run_begin = (
        paragraph.add_run()
    )

    fld_begin = OxmlElement(
        "w:fldChar"
    )

    fld_begin.set(
        qn("w:fldCharType"),
        "begin"
    )

    fld_begin.set(
        qn("w:dirty"),
        "true"
    )

    run_begin._r.append(
        fld_begin
    )

    set_page_number_font(
        run_begin
    )

    # -----------------------------
    # PAGE指令
    # -----------------------------

    run_instr = (
        paragraph.add_run()
    )

    instr = OxmlElement(
        "w:instrText"
    )

    instr.set(
        qn("xml:space"),
        "preserve"
    )

    instr.text = (
        " PAGE "
    )

    run_instr._r.append(
        instr
    )

    set_page_number_font(
        run_instr
    )

    # -----------------------------
    # separate
    # -----------------------------

    run_sep = (
        paragraph.add_run()
    )

    fld_sep = OxmlElement(
        "w:fldChar"
    )

    fld_sep.set(
        qn("w:fldCharType"),
        "separate"
    )

    run_sep._r.append(
        fld_sep
    )

    set_page_number_font(
        run_sep
    )

    # -----------------------------
    # end
    # -----------------------------

    run_end = (
        paragraph.add_run()
    )

    fld_end = OxmlElement(
        "w:fldChar"
    )

    fld_end.set(
        qn("w:fldCharType"),
        "end"
    )

    run_end._r.append(
        fld_end
    )

    set_page_number_font(
        run_end
    )


# =====================================================
# 页码格式解析
# =====================================================


def split_page_format():
    """
    根据:

        — {page} —

    分离:

        prefix

        suffix

    """

    if "{page}" not in PAGE_NUMBER_FORMAT:

        return "", ""

    prefix, suffix = (
        PAGE_NUMBER_FORMAT.split(
            "{page}",
            1
        )
    )

    return prefix, suffix


# =====================================================
# 对齐转换
# =====================================================


def get_alignment():
    """
    config字符串转换
    """

    mapping = {


        "center":
            WD_ALIGN_PARAGRAPH.CENTER,


        "left":
            WD_ALIGN_PARAGRAPH.LEFT,


        "right":
            WD_ALIGN_PARAGRAPH.RIGHT

    }

    return mapping.get(
        PAGE_NUMBER_ALIGNMENT,
        WD_ALIGN_PARAGRAPH.CENTER
    )


# =====================================================
# 添加页码
# =====================================================


def add_page_number(
        doc: Document
):
    """
    添加页码

    奇数页：
        右对齐

    偶数页：
        左对齐

    页码格式保持：

        — 1 —

    """

    prefix, suffix = split_page_format()

    for section in doc.sections:

        # ============================================
        # 奇数页页脚
        # ============================================

        odd_footer = section.footer

        clear_footer(
            odd_footer
        )

        odd_paragraph = odd_footer.paragraphs[0]

        odd_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.RIGHT
        )

        # 左侧字符

        if prefix:

            run = odd_paragraph.add_run(
                prefix
            )

            set_page_number_font(
                run
            )

        # PAGE域

        add_page_field(
            odd_paragraph
        )

        # 右侧字符

        if suffix:

            run = odd_paragraph.add_run(
                suffix
            )

            set_page_number_font(
                run
            )

        # ============================================
        # 偶数页页脚
        # ============================================

        even_footer = section.even_page_footer

        clear_footer(
            even_footer
        )

        even_paragraph = even_footer.paragraphs[0]

        even_paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        # 左侧字符

        if prefix:

            run = even_paragraph.add_run(
                prefix
            )

            set_page_number_font(
                run
            )

        # PAGE域

        add_page_field(
            even_paragraph
        )

        # 右侧字符

        if suffix:

            run = even_paragraph.add_run(
                suffix
            )

            set_page_number_font(
                run
            )

    return doc

# =====================================================
# 测试
# =====================================================


if __name__ == "__main__":

    from config import INPUT_FILE

    doc = Document(
        INPUT_FILE
    )

    add_page_number(
        doc
    )

    doc.save(
        "page_number_test.docx"
    )

    print(
        "页码生成完成"
    )
