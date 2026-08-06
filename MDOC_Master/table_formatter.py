"""
table_formatter.py

公文表格格式处理模块


功能:

1. 清除原表格所有格式

2. 表头使用一级标题样式

3. 表格正文使用正文样式

4. 设置黑色单线全部框线

5. 设置文字水平、垂直居中

6. 删除Word主题样式

7. 防止双层边框


格式依据:

中国军队机关公文格式

"""


from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from config import (
    BODY_FONT_CN,
    BODY_FONT_EN,
    BODY_FONT_SIZE,
    HEADING1_FONT_CN,
    HEADING1_FONT_EN,
    HEADING1_FONT_SIZE,
)

# =====================================================
# 清除表格原始格式
# =====================================================


def clear_table_style(table):

    tblPr = table._tbl.tblPr

    # 删除表格级属性

    remove_items = [

        "tblStyle",

        "tblLook",

        "tblBorders",

        "tblCellSpacing",

        "tblLayout",

    ]

    for item in remove_items:

        element = tblPr.find(
            qn(
                "w:" + item
            )
        )

        if element is not None:

            tblPr.remove(
                element
            )

    # 删除主题格式

    for element in tblPr.findall(
        qn("w:tblStylePr")
    ):

        tblPr.remove(
            element
        )

    # 单元格处理

    for row in table.rows:

        for cell in row.cells:

            tcPr = (
                cell._tc
                .get_or_add_tcPr()
            )

            for item in [

                "tcBorders",

                "shd",

                "tcMar",

            ]:

                element = tcPr.find(
                    qn(
                        "w:" + item
                    )
                )

                if element is not None:

                    tcPr.remove(
                        element
                    )

            # 清除段落格式

            for paragraph in cell.paragraphs:

                paragraph.alignment = None

                pf = (
                    paragraph
                    .paragraph_format
                )

                pf.space_before = Pt(0)

                pf.space_after = Pt(0)

                pf.first_line_indent = Pt(0)

                pf.left_indent = Pt(0)

                pf.right_indent = Pt(0)


# =====================================================
# 设置Run字体
# =====================================================


def set_run_font(
        run,
        cn,
        en,
        size
):

    # 清除文字效果

    run.bold = False

    run.italic = False

    run.underline = False

    run.font.strike = False

    run.font.highlight_color = None

    rPr = (
        run._element
        .get_or_add_rPr()
    )

    rFonts = (
        rPr
        .get_or_add_rFonts()
    )

    rFonts.set(
        qn("w:eastAsia"),
        cn
    )

    rFonts.set(
        qn("w:ascii"),
        en
    )

    rFonts.set(
        qn("w:hAnsi"),
        en
    )

    run.font.name = en

    run.font.size = Pt(size)


# =====================================================
# 设置单元格字体
# =====================================================


def format_cell_text(
        cell,
        header=False
):

    if header:

        cn = HEADING1_FONT_CN

        en = HEADING1_FONT_EN

        size = HEADING1_FONT_SIZE

    else:

        cn = BODY_FONT_CN

        en = BODY_FONT_EN

        size = BODY_FONT_SIZE

    for paragraph in cell.paragraphs:

        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
        )

        for run in paragraph.runs:

            set_run_font(
                run,
                cn,
                en,
                size
            )


# =====================================================
# 单元格垂直居中
# =====================================================


def set_cell_vertical_center(cell):

    cell.vertical_alignment = (
        WD_CELL_VERTICAL_ALIGNMENT.CENTER
    )


# =====================================================
# 设置表格边框
# =====================================================


def set_table_border(table):

    tblPr = table._tbl.tblPr

    # 删除旧边框

    old = tblPr.find(
        qn("w:tblBorders")
    )

    if old is not None:

        tblPr.remove(
            old
        )

    borders = OxmlElement(
        "w:tblBorders"
    )

    for edge in (

        "top",

        "left",

        "bottom",

        "right",

        "insideH",

        "insideV",

    ):

        element = OxmlElement(
            "w:" + edge
        )

        element.set(
            qn("w:val"),
            "single"
        )

        element.set(
            qn("w:sz"),
            "8"
        )

        element.set(
            qn("w:space"),
            "0"
        )

        element.set(
            qn("w:color"),
            "000000"
        )

        borders.append(
            element
        )

    tblPr.append(
        borders
    )


# =====================================================
# 设置表头重复
# =====================================================


def set_repeat_table_header(row):

    trPr = row._tr.get_or_add_trPr()

    tblHeader = OxmlElement(
        "w:tblHeader"
    )

    tblHeader.set(
        qn("w:val"),
        "true"
    )

    trPr.append(
        tblHeader
    )


# =====================================================
# 设置单个表格
# =====================================================


def format_table(table):

    # 1 清除原格式

    clear_table_style(
        table
    )

    # 2 表格居中

    table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )

    # 3 边框

    set_table_border(
        table
    )

    # 4 处理行

    for row_index, row in enumerate(
            table.rows
    ):

        # 第一行作为表头

        if row_index == 0:

            set_repeat_table_header(
                row
            )

        for cell in row.cells:

            set_cell_vertical_center(
                cell
            )

            if row_index == 0:

                format_cell_text(
                    cell,
                    header=True
                )

            else:

                format_cell_text(
                    cell,
                    header=False
                )


# =====================================================
# 总入口
# =====================================================


def format_tables(document):

    for table in document.tables:

        format_table(
            table
        )

    return document


# =====================================================
# 测试
# =====================================================

if __name__ == "__main__":

    from docx import Document

    from config import INPUT_FILE

    doc = Document(
        INPUT_FILE
    )

    format_tables(
        doc
    )

    doc.save(
        "table_test.docx"
    )

    print(
        "表格格式处理完成"
    )
