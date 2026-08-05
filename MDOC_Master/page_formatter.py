"""
page_formatter.py

Word页面版式处理模块

功能:

1. 设置A4纸张
2. 设置页面方向
3. 设置页边距
4. 设置页眉页脚距离
5. 开启奇偶页页眉页脚
6. 根据config删除原页眉
7. 根据config删除原页脚


说明:

本模块只负责页面版式。

不负责：

1. 字体
2. 标题
3. 段落
4. 页码内容


依赖:

python-docx
"""


from docx import Document

from docx.shared import Cm

from docx.enum.section import WD_ORIENT

from docx.oxml import OxmlElement

from docx.oxml.ns import qn


from config import (

    PAGE_WIDTH,
    PAGE_HEIGHT,

    TOP_MARGIN,
    BOTTOM_MARGIN,
    LEFT_MARGIN,
    RIGHT_MARGIN,

    HEADER_DISTANCE,
    FOOTER_DISTANCE,

    REMOVE_EXISTING_HEADER,
    REMOVE_EXISTING_FOOTER

)


# =====================================================
# 开启奇偶页
# =====================================================

def enable_even_odd_pages(
        document
):
    """
    开启Word奇偶页不同页眉页脚

    Word XML:

        w:evenAndOddHeaders

    """

    settings = (
        document.settings.element
    )

    flag = settings.find(
        qn(
            "w:evenAndOddHeaders"
        )
    )

    if flag is None:

        flag = OxmlElement(
            "w:evenAndOddHeaders"
        )

        settings.append(
            flag
        )


# =====================================================
# 页面方向
# =====================================================

def set_page_orientation(
        section
):

    section.orientation = (
        WD_ORIENT.PORTRAIT
    )


# =====================================================
# 页面尺寸
# =====================================================

def set_page_size(
        section
):

    section.page_width = Cm(
        PAGE_WIDTH
    )

    section.page_height = Cm(
        PAGE_HEIGHT
    )


# =====================================================
# 页边距
# =====================================================

def set_page_margin(
        section
):

    section.top_margin = Cm(
        TOP_MARGIN
    )

    section.bottom_margin = Cm(
        BOTTOM_MARGIN
    )

    section.left_margin = Cm(
        LEFT_MARGIN
    )

    section.right_margin = Cm(
        RIGHT_MARGIN
    )


# =====================================================
# 页眉页脚距离
# =====================================================

def set_header_footer_distance(
        section
):

    section.header_distance = Cm(
        HEADER_DISTANCE
    )

    section.footer_distance = Cm(
        FOOTER_DISTANCE
    )


# =====================================================
# 清空页眉
# =====================================================

def clear_header(
        header
):
    """
    清除指定页眉
    """

    for paragraph in header.paragraphs:

        paragraph.clear()


# =====================================================
# 清空页脚
# =====================================================

def clear_footer(
        footer
):
    """
    清除指定页脚
    """

    for paragraph in footer.paragraphs:

        paragraph.clear()


# =====================================================
# 删除全部页眉
# =====================================================

def remove_headers(
        section
):
    """
    删除:

    1. 奇数页页眉
    2. 偶数页页眉
    3. 首页页眉

    """

    clear_header(
        section.header
    )

    clear_header(
        section.even_page_header
    )

    clear_header(
        section.first_page_header
    )


# =====================================================
# 删除全部页脚
# =====================================================

def remove_footers(
        section
):
    """
    删除:

    1. 奇数页页脚
    2. 偶数页页脚
    3. 首页页脚

    """

    clear_footer(
        section.footer
    )

    clear_footer(
        section.even_page_footer
    )

    clear_footer(
        section.first_page_footer
    )


# =====================================================
# Section处理
# =====================================================

def format_section(
        section
):

    set_page_orientation(
        section
    )

    set_page_size(
        section
    )

    set_page_margin(
        section
    )

    set_header_footer_distance(
        section
    )

    if REMOVE_EXISTING_HEADER:

        remove_headers(
            section
        )

    if REMOVE_EXISTING_FOOTER:

        remove_footers(
            section
        )


# =====================================================
# 主接口
# =====================================================

def format_page(
        document: Document
):
    """
    页面版式统一入口
    """

    # 开启奇偶页模式

    enable_even_odd_pages(
        document
    )

    for section in document.sections:

        format_section(
            section
        )

    return document


# =====================================================
# 测试
# =====================================================
if __name__ == "__main__":

    from config import INPUT_FILE

    doc = Document(
        INPUT_FILE
    )

    format_page(
        doc
    )

    doc.save(
        "page_test.docx"
    )

    print(
        "页面版式处理完成"
    )
