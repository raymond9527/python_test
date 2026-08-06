"""
document_reader.py

Word文档读取模块

功能:
    1. 加载docx文件
    2. 检查文件有效性
    3. 获取文档基本信息
    4. 保存处理后的Word文件

依赖:
    python-docx
"""


from pathlib import Path

from docx import Document

# =====================================================
# 文件检查
# =====================================================


def check_document(
        file_path: Path
) -> bool:
    """
    检查Word文件是否有效
    """

    if not file_path.exists():

        return False

    if file_path.suffix.lower() != ".docx":

        return False

    return True


# =====================================================
# 读取Word文档
# =====================================================


def load_document(
        file_path: Path
) -> Document:
    """
    加载Word文档
    """

    if not check_document(
            file_path
    ):

        raise FileNotFoundError(
            f"无效Word文件:{file_path}"
        )

    document = Document(
        file_path
    )

    return document


# =====================================================
# 保存Word文档
# =====================================================


def save_document(
        document: Document,
        output_file: Path
):
    """
    保存Word文档


    参数:

        document:
            Document对象


        output_file:
            输出文件路径

    """

    # 创建输出目录

    output_file.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document.save(
        output_file
    )


# =====================================================
# 获取文档信息
# =====================================================


def get_document_info(
        document: Document
) -> dict:
    """
    获取Word基本信息
    """

    info = {


        "paragraphs":
            len(
                document.paragraphs
            ),


        "tables":
            len(
                document.tables
            ),


        "sections":
            len(
                document.sections
            )

    }

    return info


# =====================================================
# 清理空段落
# =====================================================


def remove_empty_paragraphs(
        document: Document
):
    """
    删除空白段落
    """

    remove_list = []

    for paragraph in document.paragraphs:

        if not paragraph.text.strip():

            remove_list.append(
                paragraph
            )

    for paragraph in remove_list:

        p = paragraph._element

        p.getparent().remove(
            p
        )


# =====================================================
# 测试接口
# =====================================================

if __name__ == "__main__":

    from config import INPUT_FILE

    print(
        "检查文件:"
    )

    print(
        INPUT_FILE
    )

    doc = load_document(
        INPUT_FILE
    )

    info = get_document_info(
        doc
    )

    print(
        "文档信息:"
    )

    for key, value in info.items():

        print(
            f"{key}: {value}"
        )
